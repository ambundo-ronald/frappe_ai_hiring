# Copyright (c) 2025, Your Company and contributors
# For license information, please see license.txt

"""
Resume Text Extraction Service
Extracts text from various resume formats (PDF, DOCX, TXT)
"""

import frappe
import os
from typing import Optional


class ResumeExtractor:
	"""Extract text from resume files"""

	def __init__(self):
		self.supported_formats = [".pdf", ".docx", ".doc", ".txt"]

	def extract_text(self, file_path: str) -> Optional[str]:
		"""
		Extract text from resume file

		Args:
			file_path: Absolute path to resume file

		Returns:
			Extracted text or None
		"""
		if not os.path.exists(file_path):
			frappe.throw(f"File not found: {file_path}")

		file_ext = os.path.splitext(file_path)[1].lower()

		if file_ext not in self.supported_formats:
			frappe.throw(f"Unsupported file format: {file_ext}")

		try:
			if file_ext == ".pdf":
				return self._extract_from_pdf(file_path)
			elif file_ext in [".docx", ".doc"]:
				return self._extract_from_docx(file_path)
			elif file_ext == ".txt":
				return self._extract_from_txt(file_path)
		except Exception as e:
			frappe.log_error(f"Failed to extract text from {file_path}: {str(e)}", "Resume Extraction")
			frappe.throw(f"Failed to extract text: {str(e)}")

		return None

	def _extract_from_pdf(self, file_path: str) -> str:
		"""Extract text from PDF file with fallback to OCR for scanned PDFs."""
		try:
			import PyPDF2

			text = []
			with open(file_path, "rb") as file:
				pdf_reader = PyPDF2.PdfReader(file)
				for page in pdf_reader.pages:
					page_text = page.extract_text()
					if page_text and page_text.strip():
						text.append(page_text)

			extracted = "\n\n".join(text)
			
			# If no text extracted (likely scanned PDF), try OCR fallback
			if not extracted or not extracted.strip():
				frappe.logger("ai_hiring").warning(f"No text extracted from PDF {file_path}, attempting OCR fallback")
				return self._extract_pdf_with_ocr(file_path)

			return extracted

		except ImportError:
			frappe.throw(
				"PyPDF2 library not installed. Install with: pip install PyPDF2"
			)
		except Exception as e:
			frappe.logger("ai_hiring").warning(f"PDF text extraction failed: {str(e)}, trying OCR fallback")
			try:
				return self._extract_pdf_with_ocr(file_path)
			except:
				raise Exception(f"PDF extraction and OCR both failed: {str(e)}")

	def _extract_from_docx(self, file_path: str) -> str:
		"""Extract text from DOCX file"""
		try:
			import docx

			doc = docx.Document(file_path)
			text = []

			for paragraph in doc.paragraphs:
				if paragraph.text.strip():
					text.append(paragraph.text)

			# Extract text from tables
			for table in doc.tables:
				for row in table.rows:
					for cell in row.cells:
						if cell.text.strip():
							text.append(cell.text)

			return "\n\n".join(text)

		except ImportError:
			frappe.throw(
				"python-docx library not installed. Install with: pip install python-docx"
			)
		except Exception as e:
			raise Exception(f"DOCX extraction failed: {str(e)}")

	def _extract_from_txt(self, file_path: str) -> str:
		"""Extract text from TXT file"""
		try:
			with open(file_path, "r", encoding="utf-8") as file:
				return file.read()
		except UnicodeDecodeError:
			# Try with different encoding
			with open(file_path, "r", encoding="latin-1") as file:
				return file.read()
		except Exception as e:
			raise Exception(f"TXT extraction failed: {str(e)}")

	def _extract_pdf_with_ocr(self, file_path: str) -> str:
		"""Extract text from PDF using OCR (for scanned/image-only PDFs)."""
		try:
			import pytesseract
			from pdf2image import convert_from_path

			frappe.logger("ai_hiring").info(f"Using OCR to extract text from {file_path}")
			images = convert_from_path(file_path)
			
			text = []
			for idx, image in enumerate(images):
				frappe.logger("ai_hiring").debug(f"OCR processing page {idx + 1}/{len(images)}")
				ocr_text = pytesseract.image_to_string(image)
				if ocr_text and ocr_text.strip():
					text.append(ocr_text)
			
			if not text:
				raise Exception("OCR failed to extract any text from PDF")
			
			frappe.logger("ai_hiring").info(f"OCR extraction successful: {len(text)} pages processed")
			return "\n\n".join(text)
		
		except ImportError:
			raise Exception(
				"OCR libraries not installed. Install with: pip install pytesseract pdf2image. "
				"Also install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki"
			)
		except Exception as e:
			raise Exception(f"OCR extraction failed: {str(e)}")

	def extract_from_attachment(self, file_url: str) -> Optional[str]:
		"""
		Extract text from Frappe file attachment

		Args:
			file_url: File URL from Frappe File doctype

		Returns:
			Extracted text or None
		"""
		try:
			# Get file path from Frappe
			file_doc = frappe.get_doc("File", {"file_url": file_url})
			
			if not file_doc:
				frappe.throw(f"File not found: {file_url}")

			# Get absolute file path
			from frappe.utils import get_files_path
			
			file_path = file_doc.get_full_path()
			
			if not file_path or not os.path.exists(file_path):
				frappe.throw(f"Physical file not found: {file_url}")

			return self.extract_text(file_path)

		except Exception as e:
			frappe.log_error(f"Failed to extract from attachment {file_url}: {str(e)}", "Resume Extraction")
			raise


def extract_resume_text(applicant_name: str) -> Optional[str]:
	"""
	Extract resume text from Job Applicant

	Args:
		applicant_name: Job Applicant name

	Returns:
		Extracted resume text or None
	"""
	applicant = frappe.get_doc("Job Applicant", applicant_name)

	if not applicant.resume_attachment:
		return None

	extractor = ResumeExtractor()
	return extractor.extract_from_attachment(applicant.resume_attachment)


@frappe.whitelist()
def test_extraction(file_url: str):
	"""Test resume extraction (for debugging)"""
	try:
		extractor = ResumeExtractor()
		text = extractor.extract_from_attachment(file_url)
		if text:
			return {"success": True, "text": text[:500], "length": len(text)}
		else:
			return {"success": False, "error": "No text extracted"}
	except Exception as e:
		return {"success": False, "error": str(e)}


@frappe.whitelist()
def debug_resume_extraction(applicant_name: str):
	"""Debug resume extraction step-by-step for troubleshooting."""
	try:
		# Step 1: Validate applicant
		if not frappe.db.exists("Job Applicant", applicant_name):
			return {"success": False, "step": 1, "error": f"Job Applicant not found: {applicant_name}"}
		
		applicant = frappe.get_doc("Job Applicant", applicant_name)
		
		# Step 2: Check resume attached
		if not applicant.resume_attachment:
			return {"success": False, "step": 2, "error": "No resume attached to applicant"}
		
		file_url = applicant.resume_attachment
		
		# Step 3: Get File document
		try:
			file_doc = frappe.get_doc("File", {"file_url": file_url})
		except:
			return {"success": False, "step": 3, "error": f"File document not found for URL: {file_url}"}
		
		# Step 4: Get file path
		try:
			file_path = file_doc.get_full_path()
		except Exception as e:
			return {"success": False, "step": 4, "error": f"Failed to get file path: {str(e)}"}
		
		# Step 5: Check file exists on disk
		import os
		if not file_path or not os.path.exists(file_path):
			return {
				"success": False,
				"step": 5,
				"error": f"Physical file not found at path: {file_path}",
				"file_path": file_path
			}
		
		# Step 6: Attempt extraction
		try:
			extractor = ResumeExtractor()
			text = extractor.extract_text(file_path)
		except Exception as e:
			return {
				"success": False,
				"step": 6,
				"error": f"Extraction failed: {str(e)}",
				"file_path": file_path,
				"file_exists": True
			}
		
		# Success
		return {
			"success": True,
			"step": "complete",
			"applicant": applicant_name,
			"file_url": file_url,
			"file_path": file_path,
			"file_exists": os.path.exists(file_path),
			"text_length": len(text) if text else 0,
			"preview": text[:300] if text else "No text extracted"
		}
	
	except Exception as e:
		return {"success": False, "error": f"Unexpected error: {str(e)}", "traceback": frappe.get_traceback()}
